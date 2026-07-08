"""Text and image extraction from documents (.txt, .md, .docx, .pdf) and standalone images.

Supports:
- Plain text files (.txt, .md)
- Word documents (.docx) via python-docx
- PDF documents (.pdf) via PyMuPDF with optional Tesseract OCR
- Standalone image files (.png, .jpg, .jpeg, .tiff, .bmp, .webp) via PIL + pytesseract

No imports from core/ — self-contained backend module.
"""

from __future__ import annotations

import io
import logging
import os
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# Optional dependencies — graceful fallback
try:
    import docx
except ImportError:
    docx = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import pytesseract
    from PIL import Image
    pytesseract.get_tesseract_version()
    OCR_AVAILABLE = True
except Exception:
    pytesseract = None
    Image = None
    OCR_AVAILABLE = False


IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"}

# ── Internal extractors ──────────────────────────────────────────────────────


def _extract_plain(file_path: str) -> str:
    """Extract text from plain text files (.txt, .md)."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def _extract_docx(file_path: str, log_callback: Optional[callable] = None) -> str:
    """Extract text from Word documents (.docx) via python-docx."""
    if docx is None:
        raise ImportError(
            "python-docx is required for .docx files. "
            "Install with: pip install python-docx"
        )

    try:
        document = docx.Document(file_path)
        paragraphs = [p.text for p in document.paragraphs]
        text = "\n".join(paragraphs)

        if not text.strip():
            # Try extracting from tables as fallback
            table_text = []
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
            if table_text:
                text = "\n".join(table_text)

        return text
    except Exception as e:
        raise Exception(
            f"Failed to read Word document. "
            f"If it's an old .doc format, convert to .docx first. Error: {e}"
        )


def _extract_pdf(file_path: str, log_callback: Optional[callable] = None, ocr_languages: str = "rus+eng") -> str:
    """Extract text from PDF files via PyMuPDF, with optional OCR fallback."""
    if fitz is None:
        raise ImportError(
            "PyMuPDF is required for .pdf files. "
            "Install with: pip install PyMuPDF"
        )

    try:
        doc = fitz.open(file_path)
        pages_text = []
        ocr_used = False

        for page_num, page in enumerate(doc):
            text = page.get_text()

            if text.strip():
                pages_text.append(text)
            elif OCR_AVAILABLE:
                if log_callback:
                    log_callback(f"  Page {page_num + 1}: no text found, running OCR...")
                ocr_text = _ocr_pdf_page(page, lang=ocr_languages)
                if ocr_text:
                    pages_text.append(ocr_text)
                    ocr_used = True
                else:
                    if log_callback:
                        log_callback(f"  ⚠️ Page {page_num + 1}: OCR returned no text")
            else:
                if log_callback:
                    log_callback(
                        f"  ⚠️ Page {page_num + 1}: no text found "
                        f"(install Tesseract-OCR for scanned PDFs)"
                    )

        doc.close()

        if ocr_used and log_callback:
            log_callback("✅ OCR recognition complete.")

        return "\n".join(pages_text)
    except Exception as e:
        raise Exception(f"Failed to read PDF file: {e}")


def _ocr_pdf_page(page, lang: str = "rus+eng") -> str:
    """Recognize text on a PDF page via Tesseract OCR."""
    if not OCR_AVAILABLE:
        return ""

    try:
        pix = page.get_pixmap(dpi=150)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        text = pytesseract.image_to_string(img, lang=lang)
        return text.strip()
    except Exception as e:
        logger.warning("OCR failed on PDF page: %s", e)
        return ""


def _extract_image(file_path: str, log_callback: Optional[callable] = None, ocr_languages: str = "rus+eng") -> str:
    """Extract text from a standalone image file via Tesseract OCR.

    Args:
        file_path: Path to the image file (.png, .jpg, .jpeg, .tiff, .bmp, .webp).
        log_callback: Optional logging callback.
        ocr_languages: Tesseract language string (e.g., "rus+eng"). Default: "rus+eng".

    Returns:
        Extracted text content.

    Raises:
        ValueError: If pytesseract is not available.
    """
    if not OCR_AVAILABLE:
        raise ImportError(
            "Tesseract-OCR is required for image files. "
            "Install Tesseract and python packages: pip install pytesseract pillow"
        )

    try:
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img, lang=ocr_languages)
        return text.strip()
    except Exception as e:
        if log_callback:
            log_callback(f"⚠️ Failed to OCR image: {e}")
        logger.warning("OCR failed on image %s: %s", file_path, e)
        return ""


def extract_text(file_path: str | Path, log_callback: Optional[callable] = None, ocr_languages: str = "rus+eng") -> str:
    """Extract text content from a file.

    Auto-detects file type by extension and delegates to the appropriate extractor.

    Args:
        file_path: Path to the file.
        log_callback: Optional logging callback function.
        ocr_languages: Tesseract language string (e.g., "rus+eng"). Used for PDF OCR
                       and standalone image files. Default: "rus+eng".

    Returns:
        Extracted text content.

    Raises:
        FileNotFoundError: If file does not exist.
        ValueError: If file format is not supported.
        Exception: If extraction fails.
    """
    file_path = str(file_path)
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = Path(file_path).suffix.lower()

    if ext in (".txt", ".md"):
        return _extract_plain(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file_path, log_callback=log_callback)
    elif ext == ".pdf":
        return _extract_pdf(file_path, log_callback=log_callback, ocr_languages=ocr_languages)
    elif ext in IMAGE_EXTENSIONS:
        return _extract_image(file_path, log_callback=log_callback, ocr_languages=ocr_languages)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
